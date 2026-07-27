#!/usr/bin/env python3
"""生成用于 Common Agent RAGFlow 召回质量验收的多页 PDF 和 DOCX。

运行方式：
    uv run --with python-docx --with reportlab scripts/generate-knowledge-samples.py

生成内容与 corpus 目录下的 md/txt 语料同属一套虚构民生资料，
覆盖 PDF 与 Word 两种二进制格式的解析链路。
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer


REPOSITORY_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = REPOSITORY_ROOT / "test-data" / "knowledge-base" / "corpus"
FONT_NAME = "Sarasa UI SC"

PDF_FILE_NAME = "03-medical-insurance-guide.pdf"
DOCX_FILE_NAME = "04-housing-fund-guide.docx"

PDF_TITLE = "青川市城乡居民基本医疗保险报销指南"
PDF_SUBTITLE = "文件编号 QC-YB-2027-01 | 2027 年 1 月 1 日起施行 | 虚构测试资料"
PDF_FOOTER = "青川市城乡居民医保报销指南 | 虚构测试资料"

DOCX_TITLE = "青川市住房公积金提取与贷款办事指南"
DOCX_SUBTITLE = "文件编号 QC-GJJ-2027-02 | 2027 年 1 月 1 日起施行 | 虚构测试资料"
DOCX_HEADER = "青川市住房公积金办事指南 | 虚构测试资料"


MEDICAL_SECTIONS: tuple[tuple[str, tuple[str, ...]], ...] = (
    (
        "一、参保缴费",
        (
            "城乡居民医保按年度参保，集中缴费期为每年 9 月 1 日至 12 月 31 日，"
            "次年 1 月 1 日至 12 月 31 日享受待遇。错过集中缴费期的，可在次年 "
            "1 月 1 日至 2 月底补缴，补缴后有 90 天等待期，等待期内发生的医疗费用不予报销。",
            "2027 年度个人缴费标准为每人每年 420 元，财政补助每人每年 700 元。"
            "特困人员、低保对象、重度残疾人由政府全额代缴，本人不缴费；"
            "低保边缘家庭成员减半缴费，个人只需缴 210 元。",
            "新生儿在出生后 90 天内办理参保登记并缴费的，自出生之日起享受待遇，"
            "不设等待期。超过 90 天办理的，按补缴处理并适用 90 天等待期。"
            "请注意区分：新生儿宽限期是 90 天，补缴等待期也是 90 天，但两者含义完全不同。",
        ),
    ),
    (
        "二、门诊待遇",
        (
            "普通门诊起付线为每人每年 200 元，年度累计计算。起付线以上部分，"
            "在社区卫生服务中心（乡镇卫生院）就诊报销 70%，在二级及以上医院就诊报销 60%。"
            "普通门诊年度最高支付限额为 1500 元。",
            "2026 年度普通门诊起付线曾为 150 元，该标准自 2027 年 1 月 1 日起调整为 200 元。"
            "查询报销政策时应当采用现行的 200 元，150 元已经废止；"
            "如有人按旧标准询问，应说明该值已不再适用。",
            "门诊慢特病（高血压、糖尿病、恶性肿瘤门诊治疗、尿毒症透析等 28 个病种）"
            "单独计算，不占用普通门诊限额。慢特病需经二级及以上定点医院认定，"
            "认定材料受理后 15 个工作日内出具结果。",
        ),
    ),
    (
        "三、住院待遇",
        (
            "住院起付线按医院等级划分：一级医院 300 元、二级医院 600 元、三级医院 1000 元。"
            "同一年度内第二次及以后住院，起付线减半。转诊住院的，起付线只计算最高一级医院一次。",
            "住院报销比例同样按医院等级划分：一级医院 85%、二级医院 75%、三级医院 60%。"
            "起付线和报销比例是两组不同的数字，一级医院是起付线最低、报销比例最高，"
            "三级医院则相反，回答时不要把两组数字对调。",
            "基本医保年度最高支付限额为 25 万元。超过基本医保限额的部分进入大病保险，"
            "大病保险起付线为 1.5 万元，起付线以上 10 万元以内报销 60%，"
            "10 万元以上部分报销 70%，大病保险年度最高支付限额 40 万元。",
        ),
    ),
    (
        "四、异地就医",
        (
            "异地长期居住、异地转诊、异地急诊抢救三类人员可以办理异地就医备案。"
            "备案可通过国家医保服务平台 App、青川市医保小程序或参保地经办窗口办理，"
            "线上备案即时生效，窗口备案 2 个工作日内生效。",
            "已备案的异地就医按参保地同级医院标准报销，比例不变。"
            "未备案自行到异地就医的，报销比例在原基础上下降 20 个百分点，"
            "例如三级医院原本报销 60%，未备案则只报销 40%。急诊抢救可在入院后 5 个工作日内补办备案。",
            "异地就医已实现直接结算的，出院时只需支付个人负担部分，无需垫付后回参保地报销。"
            "未能直接结算的，须保留发票原件、费用清单、出院小结和医保电子凭证记录，回参保地手工报销。",
        ),
    ),
    (
        "五、报销办理与时限",
        (
            "手工报销申请应当自出院之日起 12 个月内提出，逾期不再受理。"
            "门诊费用手工报销的，应当自结算之日起 12 个月内提出。"
            "请注意这里是 12 个月，与住房公积金购房提取的 2 年期限、快递理赔的 1 年期限都不相同。",
            "手工报销材料齐全的，经办机构自受理之日起 30 个工作日内完成审核并支付。"
            "材料不全的，应当一次性告知需要补正的全部内容，补正材料之日重新起算受理时间。"
            "30 个工作日是医保手工报销的时限，不要与物业投诉 3 个工作日、"
            "公积金提取 3 个工作日等其他时限混用。",
            "不予报销的情形包括：工伤、交通事故等第三方责任的医疗费用，"
            "境外就医费用，美容整形、健康体检、预防性疫苗接种费用，"
            "以及非定点医疗机构就医费用（急诊抢救除外）。",
        ),
    ),
    (
        "六、咨询与监督",
        (
            "医保政策咨询电话 12393，全市统一。业务经办网点为各区（县）医保服务大厅，"
            "以及和风苑社区所在街道的便民服务中心医保窗口，窗口办公时间为工作日 9:00—17:00，"
            "午间不休息。",
            "参保人对报销结果有异议的，可在收到结算单之日起 30 日内申请复核，"
            "经办机构 15 个工作日内答复。对复核结果仍有异议的，可申请行政复议或者提起行政诉讼。",
            "举报欺诈骗保行为的，经查证属实按查实金额的 2% 给予奖励，"
            "单次奖励最高 20 万元，最低 200 元。举报电话与咨询电话相同，均为 12393。",
        ),
    ),
)


HOUSING_FUND_SECTIONS: tuple[tuple[str, tuple[str, ...]], ...] = (
    (
        "一、缴存",
        (
            "单位和职工的住房公积金缴存比例均不得低于 5%、不得高于 12%，"
            "单位与个人执行相同比例。缴存比例每年可以调整一次，调整时间为每年 7 月，"
            "调整后从当年 7 月 1 日起执行。",
            "缴存基数为职工本人上一年度月平均工资，上限为青川市上一年度职工月平均社会平均工资的 3 倍，"
            "2027 年度上限为 30600 元；下限为本市最低工资标准，2027 年度为 2360 元。"
            "缴存基数每年 7 月调整一次，与缴存比例的调整时间一致。",
            "新设立单位应当自设立之日起 30 日内办理缴存登记，并于登记之日起 20 日内为职工办理账户设立。"
            "新招用职工应当自用工之日起 30 日内办理缴存登记。逾期办理的，由公积金管理中心责令限期办理，"
            "逾期仍不办理的可申请人民法院强制执行。",
        ),
    ),
    (
        "二、提取",
        (
            "可以提取住房公积金的情形包括：购买、建造、翻建、大修自住住房；偿还自住住房贷款本息；"
            "租赁自住住房；离休、退休；完全丧失劳动能力并与单位终止劳动关系；"
            "出境定居；本人或配偶、父母、子女患重大疾病造成家庭生活严重困难。",
            "购房提取应当自购房之日起 2 年内首次申请，逾期不再受理首次提取。"
            "这里的购房之日，商品房以购房合同网签备案日期为准，二手房以不动产权证登记日期为准。"
            "注意这个 2 年期限与医保手工报销的 12 个月、养犬登记的 30 日都不相同。",
            "租房提取每年办理一次，无需提供租赁发票，按定额提取，"
            "每月提取限额 1200 元、每年最高 14400 元。2026 年度租房提取年度限额为 12000 元，"
            "该标准自 2027 年 1 月 1 日起提高至 14400 元，12000 元已经废止。",
        ),
    ),
    (
        "三、贷款",
        (
            "申请公积金贷款须连续足额缴存满 6 个月。贷款额度上限为：单人申请 60 万元，"
            "夫妻双方共同申请 100 万元。同时不得超过账户余额的 15 倍，也不得超过购房总价扣除首付款后的金额。"
            "三个条件取最低值。",
            "首套房最低首付比例 20%，二套房最低首付比例 30%，第三套及以上不予贷款。"
            "贷款期限最长 30 年，且借款人年龄与贷款期限之和不得超过法定退休年龄后 5 年。",
            "2027 年公积金贷款利率：五年以下（含）年利率 2.35%，五年以上年利率 2.85%。"
            "二套房贷款利率按同期首套房利率的 1.1 倍执行。提前还款不收违约金，"
            "但需在还款日前 15 个工作日提出申请。",
        ),
    ),
    (
        "四、办理时限与渠道",
        (
            "提取业务材料齐全的，自受理之日起 3 个工作日内审核完毕并划转到账。"
            "贷款业务自受理之日起 10 个工作日内完成审批。"
            "提取的 3 个工作日与《和风苑社区物业服务规范》中的投诉答复 3 个工作日数值相同但事项无关，"
            "回答时须明确说明是哪一项业务的时限。",
            "线上办理渠道为青川市住房公积金 App 与市政务服务网，"
            "线下办理渠道为各区公积金管理部服务大厅，办公时间为工作日 9:00—17:00。"
            "线上可办理的业务包括：租房提取、离退休提取、账户查询、还贷提取签约。"
            "购房提取与重大疾病提取须到窗口办理。",
            "咨询电话 12329，全市统一。对办理结果有异议的，可在收到告知书之日起 30 日内申请复核，"
            "管理中心 15 个工作日内答复。",
        ),
    ),
)


def _set_docx_run(
    run, *, size: float, bold: bool = False, color: str = "222222"
) -> None:
    run.font.name = FONT_NAME
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("eastAsia", "ascii", "hAnsi"):
        fonts.set(qn(f"w:{key}"), FONT_NAME)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def _configure_docx(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.9)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    normal = document.styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.25
    for style_name, size, before, after, color in (
        ("Heading 1", 16, 18, 10, "2E74B5"),
        ("Heading 2", 13, 14, 7, "1F4D78"),
    ):
        style = document.styles[style_name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_docx_run(header.add_run(DOCX_HEADER), size=9, color="666666")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_docx_run(
        footer.add_run("Common Agent Knowledge Fixture"), size=9, color="777777"
    )


def generate_docx() -> None:
    document = Document()
    _configure_docx(document)
    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(5)
    _set_docx_run(title.add_run(DOCX_TITLE), size=23, bold=True, color="0B2545")
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    _set_docx_run(subtitle.add_run(DOCX_SUBTITLE), size=10.5, color="555555")

    for index, (heading, paragraphs) in enumerate(HOUSING_FUND_SECTIONS):
        if index:
            document.add_page_break()
        document.add_heading(heading, level=1)
        for content in paragraphs:
            document.add_paragraph(content)
        if index == 1:
            document.add_heading("提取材料清单", level=2)
            document.add_paragraph(
                "购房提取：身份证、购房合同或不动产权证、全额购房发票、本人银行卡。"
                "租房提取：身份证、本人银行卡，无需租赁合同与发票。"
                "重大疾病提取：身份证、二级以上医院诊断证明、医保结算单、亲属关系证明。"
            )
        if index == 2:
            document.add_heading("不予贷款的情形", level=2)
            document.add_paragraph(
                "已有两笔公积金贷款未结清的、缴存不足 6 个月的、"
                "购买非自住商业用房的、征信存在连续三期或累计六期逾期的，均不予贷款。"
                "商转公业务暂不受理。"
            )

    document.core_properties.title = f"{DOCX_TITLE}（虚构测试资料）"
    document.core_properties.subject = "Common Agent RAGFlow DOCX 召回测试"
    document.core_properties.author = "Common Agent Test Fixture"
    document.core_properties.keywords = "fictional,ragflow,retrieval,housing-fund"
    document.save(OUTPUT_DIR / DOCX_FILE_NAME)


def _pdf_page(canvas, document) -> None:
    canvas.saveState()
    canvas.setFont("STSong-Light", 8.5)
    canvas.setFillColor(HexColor("#777777"))
    canvas.drawString(inch, 0.45 * inch, PDF_FOOTER)
    canvas.drawRightString(7.5 * inch, 0.45 * inch, f"第 {document.page} 页")
    canvas.restoreState()


def generate_pdf() -> None:
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    base = getSampleStyleSheet()
    title = ParagraphStyle(
        "ZhTitle",
        parent=base["Title"],
        fontName="STSong-Light",
        fontSize=21,
        leading=30,
        textColor=HexColor("#0B2545"),
        alignment=TA_CENTER,
        spaceAfter=8,
    )
    subtitle = ParagraphStyle(
        "ZhSubtitle",
        parent=base["Normal"],
        fontName="STSong-Light",
        fontSize=10,
        leading=16,
        textColor=HexColor("#666666"),
        alignment=TA_CENTER,
        spaceAfter=20,
    )
    heading = ParagraphStyle(
        "ZhHeading",
        parent=base["Heading2"],
        fontName="STSong-Light",
        fontSize=15,
        leading=22,
        textColor=HexColor("#2E74B5"),
        alignment=TA_LEFT,
        spaceAfter=10,
    )
    body = ParagraphStyle(
        "ZhBody",
        parent=base["BodyText"],
        fontName="STSong-Light",
        fontSize=11,
        leading=20,
        firstLineIndent=22,
        textColor=HexColor("#222222"),
        alignment=TA_LEFT,
        spaceAfter=12,
    )
    pdf = SimpleDocTemplate(
        str(OUTPUT_DIR / PDF_FILE_NAME),
        pagesize=letter,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=0.8 * inch,
        bottomMargin=0.7 * inch,
        title=f"{PDF_TITLE}（虚构测试资料）",
        author="Common Agent Test Fixture",
        subject="Common Agent RAGFlow PDF 召回测试",
    )
    story = [
        Paragraph(PDF_TITLE, title),
        Paragraph(PDF_SUBTITLE, subtitle),
    ]
    for index, (section_title, paragraphs) in enumerate(MEDICAL_SECTIONS):
        if index:
            story.append(PageBreak())
        story.append(Paragraph(section_title, heading))
        for content in paragraphs:
            story.append(Paragraph(content, body))
        story.append(Spacer(1, 6))
        story.append(
            Paragraph(
                "本页数值以 2027 年 1 月 1 日起施行的标准为准，此前公布的旧标准一律废止。",
                subtitle,
            )
        )
    pdf.build(story, onFirstPage=_pdf_page, onLaterPages=_pdf_page)


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    generate_pdf()
    generate_docx()
    print(f"RAGFlow 召回质量样本已生成：{OUTPUT_DIR}")


if __name__ == "__main__":
    main()
